import sys
from easygoogletranslate import EasyGoogleTranslate
import os
import PyPDF2
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
import docx
import dotenv

dotenv.load_dotenv()
def split_text(text, max_length=4000):
    chunks = []
    while len(text) > max_length:
        split_at = text.rfind(' ', 0, max_length)
        if split_at == -1:
            split_at = max_length
        chunks.append(text[:split_at])
        text = text[split_at:].strip()
    chunks.append(text)
    return chunks

# Function to extract text from image using OCR
def extract_text_from_image(image_path):
    # Open the image file
    img = Image.open(image_path)

    # Convert image to grayscale
    gray_img = img.convert('L')

    # Perform OCR to extract text
    text = pytesseract.image_to_string(gray_img)

    translator = EasyGoogleTranslate(
        source_language='en',
        target_language='gom',
        timeout=30
    )

    print("Enter Sample Text in English: ")
    engsample= text

    print("\n\nTranslating Input Text...")
    chunks = split_text(engsample)
    translated_chunks = [translator.translate(chunk) for chunk in chunks]
    result = ' '.join(translated_chunks)
    document = docx.Document()
    document.add_paragraph(result)
    document.save(os.getenv('savedDocx'))
    print("\n\nTranslated Text: ")
    print(result)
    return text

# Function to extract text from PDF using PyPDF2
def extract_text_from_pdf(pdf_path):
    # Open the PDF file in read-binary mode
    with open(pdf_path, 'rb') as pdf_file:
        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Initialize list to store extracted text
        extracted_text = []

        # Iterate through each page of the PDF
        for page_num in range(len(pdf_reader.pages)):
            # Get the text from the current page
            page_text = pdf_reader.pages[page_num].extract_text()
            extracted_text.append(page_text)

        # Combine all extracted text into a single string
        pdf_text = '\n'.join(extracted_text)
        
        translator = EasyGoogleTranslate(
        source_language='en',
        target_language='gom',
        timeout=30
        )

        engsample= pdf_text
        print(engsample)
        
        print("\n\nTranslating Input Text...")
        chunks = split_text(engsample)
        translated_chunks = [translator.translate(chunk) for chunk in chunks]
        result = ' '.join(translated_chunks)
        document = docx.Document()
        document.add_paragraph(result)
        document.save(os.getenv('savedDocx'))
        print("\n\nTranslated Text: ")
        print(result)
        return pdf_text

# Function to extract text based on file type (image or PDF)
def extract_text(input_path):
    # Get the file extension
    file_ext = os.path.splitext(input_path)[1].lower()
    # Check if input is an image
    if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
        # Extract text from image using OCR
        text = extract_text_from_image(input_path)
    # Check if input is a PDF
    elif file_ext == '.pdf':
        # Extract text from PDF
        text = extract_text_from_pdf(input_path)
    else:
        text = "Unsupported file format"
    return text

if len(sys.argv) != 2:
    print("Usage: python Tesseract_OCR_test.py")
    sys.exit(1)

# Path to the uploaded file provided as command-line argument
input_path = sys.argv[1]
extract_text(input_path)